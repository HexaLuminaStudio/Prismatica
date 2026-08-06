# coding: utf-8
"""
HSK 检索结果导出 Worker(PRD-005 + 修复版)
==========================================

支持两种导出模式:
    - mergeMode=False  (默认):每篇作文一个独立 txt/docx 文件
    - mergeMode=True   (新)   :所有作文合并到一个 txt/docx 文件,篇间用分隔线

写入前自动清洗 XML 不兼容字符(NULL / 控制字符 / DEL 等),
避免 python-docx 抛 "All strings must be XML compatible" 错误。

信号:
    progress = Signal(int, int)                       # current, total
    wroteFile = Signal(str)                            # 文件写入成功
    finishedWithResult = Signal(int, int, int)         # success, skipped, fail
    failed = Signal(str)
"""

from __future__ import annotations

import re
import time
from pathlib import Path
from typing import Any, Dict, List, Literal, Optional

from PySide6.QtCore import QThread, Signal

from app.core.utils import log


# ---------------------------------------------------------------------------
# 文件名清洗
# ---------------------------------------------------------------------------
def _sanitizeFilename(text: str, maxLen: int = 20) -> str:
    if not text:
        return "untitled"
    # 1) 剥掉 HSK 偏误标注 [BD《]...[BD》] / [WD,3]1 / {CC} / 《...》等
    cleaned = re.sub(r"\[(?:BD|WD|YY|XQ|YQ)[^]]*\]", "", text)
    cleaned = re.sub(r"\{(?:CC|CQ|YY|XQ|B|CY)\}", "", cleaned)
    cleaned = re.sub(r"《[^》]*》", "", cleaned)
    # 2) 替换 Windows/Unix 非法字符
    cleaned = re.sub(r'[\\/:*?"<>|\r\n\t]+', '_', cleaned)
    cleaned = cleaned.strip().strip('.')
    if not cleaned:
        cleaned = "untitled"
    return cleaned[:maxLen]


# ---------------------------------------------------------------------------
# XML 安全文本清洗
# ---------------------------------------------------------------------------
# XML 1.0 不允许的字符:
#   - NULL (0x00)
#   - 除 \t(0x09) \n(0x0A) \r(0x0D) 之外的所有 0x00-0x1F 控制字符
#   - DEL (0x7F)
#   - 0xFFFE / 0xFFFF (BOM 字符)
_XML_FORBIDDEN_RE = re.compile(
    r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F\uFFFE\uFFFF]"
)


def _sanitizeForXml(text: str) -> str:
    """把 XML 不兼容字符替换成 '?'(避免 docx 写入失败)。"""
    if not text:
        return text or ""
    return _XML_FORBIDDEN_RE.sub("?", text)


# ---------------------------------------------------------------------------
# Worker
# ---------------------------------------------------------------------------
class HskCorpusExportWorker(QThread):
    """后台批量导出 Worker。"""

    progress = Signal(int, int)  # current, total
    wroteFile = Signal(str)
    finishedWithResult = Signal(int, int, int)  # success, skipped, fail
    failed = Signal(str)

    # 默认合并文件名(用户未指定时用这个)
    DEFAULT_MERGE_NAME = "hsk_export"

    def __init__(
        self,
        zwhaoList: List[str],
        outputDir: str,
        fileFormat: Literal["txt", "docx"],
        skipMissingTitle: bool = False,
        mergeMode: bool = False,
        mergeFileName: Optional[str] = None,
        parent=None,
    ) -> None:
        super().__init__(parent)
        self._zwhaoList = list(zwhaoList)
        self._outputDir = Path(outputDir)
        self._fileFormat = fileFormat
        # skipMissingTitle 默认 False:只要 zwhao 在 local db 就导出,
        # 不再用「是否含 Title」过滤 — 作文母号才是唯一标识。
        self._skipMissingTitle = bool(skipMissingTitle)
        self._mergeMode = bool(mergeMode)
        self._mergeFileName = (
            mergeFileName or self.DEFAULT_MERGE_NAME
        )
        self._isRunning = True

    def stop(self) -> None:
        self._isRunning = False

    # ------------------------------------------------------------------
    # 主循环
    # ------------------------------------------------------------------
    def run(self) -> None:
        from app.core.services.hsk_local_corpus_service import (
            hskLocalCorpusService,
        )

        # 0. 输出目录
        try:
            self._outputDir.mkdir(parents=True, exist_ok=True)
        except Exception as e:
            log.error(f"[HskCorpusExportWorker] 创建输出目录失败: {e}")
            self.failed.emit(f"无法创建输出目录: {e}")
            return

        # 1. local db 可用性
        if not hskLocalCorpusService.isAvailable():
            self.failed.emit(
                "本地镜像库不可用,请检查 datas/corpora/hsk_corpus_local.db"
            )
            return

        total = len(self._zwhaoList)
        if total == 0:
            log.info("[HskCorpusExportWorker] 命中列表为空,跳过")
            self.finishedWithResult.emit(0, 0, 0)
            return

        # 2. 一次性取所有记录
        modeLabel = "合并" if self._mergeMode else "分文件"
        log.info(
            f"[HskCorpusExportWorker] 开始导出: {total} 条 → "
            f"{self._outputDir}, 格式={self._fileFormat}, "
            f"模式={modeLabel}, skipMissingTitle={self._skipMissingTitle}"
        )
        t0 = time.time()
        records = hskLocalCorpusService.fetchRecordsByZwhaoList(
            self._zwhaoList
        )
        recordByZwhao: Dict[str, Dict[str, Any]] = {
            r["zwhao"]: r for r in records
        }
        log.info(
            f"[HskCorpusExportWorker] 从 local db 取出 {len(records)} 条"
        )

        # 3. 进度
        successCount = 0
        skippedCount = 0
        failCount = 0
        processedCount = 0

        # 合并模式:一个文档累积写入
        mergeDoc = None  # docx.Document 实例(合并 docx 模式用)
        mergeFilePath: Optional[Path] = None
        if self._mergeMode:
            mergeFilePath = self._outputDir / (
                f"{self._mergeFileName}.{self._fileFormat}"
            )
            if self._fileFormat == "docx":
                from docx import Document
                mergeDoc = Document()
                # 顶部总标题
                mergeDoc.add_heading(
                    f"HSK 作文合集 - 共 {total} 篇", level=0
                )
                mergeDoc.add_paragraph(
                    f"导出时间: {time.strftime('%Y-%m-%d %H:%M:%S')}"
                )
                mergeDoc.add_paragraph("")

        for zwhao in self._zwhaoList:
            if not self._isRunning:
                log.info("[HskCorpusExportWorker] 用户取消")
                break

            processedCount += 1
            rec = recordByZwhao.get(zwhao)

            if rec is None:
                # zwhao 不在 local db(11337 vs 11328 差异)
                log.warning(
                    f"[HskCorpusExportWorker] zwhao 不在 local db: {zwhao}"
                )
                skippedCount += 1
                self.progress.emit(processedCount, total)
                continue

            # 作文母号(zwhao)是唯一标识 — 只要在 local db 就导出,
            # 不再用「是否含 Title」过滤。无 Title 的作文会正常导出,
            # 元信息里写「(未提取到篇目)」。
            # (skipMissingTitle 参数保留为向后兼容,但默认 False,不再过滤)

            try:
                if self._mergeMode:
                    # 合并模式:追加写入同一文件
                    self._appendToMerge(rec, mergeDoc, mergeFilePath)
                else:
                    # 分文件模式:每篇一个文件
                    filePath = self._writeRecord(rec)
                    self.wroteFile.emit(str(filePath))
                successCount += 1
            except Exception as e:
                log.error(
                    f"[HskCorpusExportWorker] 写入 {zwhao} 失败: "
                    f"{type(e).__name__}: {e}"
                )
                failCount += 1

            self.progress.emit(processedCount, total)

        # 合并模式:关闭/保存
        if self._mergeMode and mergeDoc is not None:
            try:
                mergeDoc.save(str(mergeFilePath))
                self.wroteFile.emit(str(mergeFilePath))
                log.info(
                    f"[HskCorpusExportWorker] 合并文件已保存: {mergeFilePath}"
                )
            except Exception as e:
                log.error(
                    f"[HskCorpusExportWorker] 保存合并文件失败: {e}"
                )
                failCount += 1
                successCount = max(0, successCount - 1)
        elif self._mergeMode and self._fileFormat == "txt":
            # txt 已在 _appendToMerge 内逐条写入
            log.info(
                f"[HskCorpusExportWorker] txt 合并文件已写入: {mergeFilePath}"
            )

        elapsed = time.time() - t0
        log.info(
            f"[HskCorpusExportWorker] 完成:成功 {successCount},"
            f"跳过 {skippedCount},失败 {failCount},"
            f"耗时 {elapsed:.2f}s"
        )
        self.finishedWithResult.emit(successCount, skippedCount, failCount)

    # ------------------------------------------------------------------
    # 合并模式写入(追加到同一文件)
    # ------------------------------------------------------------------
    def _appendToMerge(
        self,
        rec: Dict[str, Any],
        mergeDoc,  # docx.Document or None
        mergeFilePath: Path,
    ) -> None:
        zwhao = rec["zwhao"]
        title = rec.get("title") or ""
        dataText = rec.get("data") or ""
        fetchedAt = rec.get("fetchedAt") or ""

        if self._fileFormat == "txt":
            content = self._formatTxtContent(
                zwhao, title, dataText, fetchedAt
            )
            # 追加写(已有内容则换行分隔)
            with open(mergeFilePath, "a", encoding="utf-8") as f:
                # XML 安全清洗
                f.write(_sanitizeForXml(content))
                f.write("\n\n" + "=" * 60 + "\n\n")
        else:  # docx
            # Heading: 标题
            heading_text = title if title else f"作文 {zwhao}"
            mergeDoc.add_heading(_sanitizeForXml(heading_text), level=1)
            # 元数据
            mergeDoc.add_paragraph(
                _sanitizeForXml(f"作文母号: {zwhao}")
            )
            if fetchedAt:
                mergeDoc.add_paragraph(
                    _sanitizeForXml(f"获取时间: {fetchedAt}")
                )
            mergeDoc.add_paragraph("")  # 空行
            # 正文按行写(每行清洗)
            for line in dataText.split("\n"):
                mergeDoc.add_paragraph(_sanitizeForXml(line))
            # 篇间分隔(用页分隔符)
            from docx.enum.text import WD_BREAK
            mergeDoc.add_paragraph().add_run().add_break(
                WD_BREAK.PAGE
            )

    # ------------------------------------------------------------------
    # 分文件模式(每篇一个文件)
    # ------------------------------------------------------------------
    def _writeRecord(self, rec: Dict[str, Any]) -> Path:
        zwhao = rec["zwhao"]
        title = rec.get("title") or ""
        dataText = rec.get("data") or ""
        fetchedAt = rec.get("fetchedAt") or ""

        safeTitle = _sanitizeFilename(title)
        ext = self._fileFormat
        fileName = f"{zwhao}_{safeTitle}.{ext}"
        filePath = self._outputDir / fileName

        if self._fileFormat == "txt":
            content = self._formatTxtContent(zwhao, title, dataText, fetchedAt)
            filePath.write_text(_sanitizeForXml(content), encoding="utf-8")
        else:  # docx
            self._writeDocx(filePath, zwhao, title, dataText, fetchedAt)

        return filePath

    # ------------------------------------------------------------------
    # 内容格式化
    # ------------------------------------------------------------------
    @staticmethod
    def _formatTxtContent(
        zwhao: str, title: str, dataText: str, fetchedAt: str
    ) -> str:
        parts = [
            f"# 作文母号: {zwhao}",
            f"# 标题: {title}" if title else "# 标题: (未提取到篇目)",
        ]
        if fetchedAt:
            parts.append(f"# 获取时间: {fetchedAt}")
        parts.append("")
        parts.append(dataText)
        return "\n".join(parts)

    @staticmethod
    def _writeDocx(
        filePath: Path,
        zwhao: str,
        title: str,
        dataText: str,
        fetchedAt: str,
    ) -> None:
        from docx import Document

        doc = Document()
        heading = title if title else f"作文 {zwhao}"
        doc.add_heading(_sanitizeForXml(heading), level=1)

        metaLines = [_sanitizeForXml(f"作文母号: {zwhao}")]
        if fetchedAt:
            metaLines.append(_sanitizeForXml(f"获取时间: {fetchedAt}"))
        for line in metaLines:
            doc.add_paragraph(line)

        doc.add_paragraph("")

        # 正文按行写入,每行清洗
        for line in dataText.split("\n"):
            doc.add_paragraph(_sanitizeForXml(line))

        doc.save(str(filePath))
