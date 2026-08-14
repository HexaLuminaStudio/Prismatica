# coding: utf-8
"""偏误分析多格式导入回归测试。"""

from __future__ import annotations

import struct

from docx import Document
from openpyxl import Workbook

from app.core.services.bias_document_service import (
    BIAS_TEXT_COLUMN,
    BiasDocumentLoadError,
    BiasDocumentService,
)
from app.view.bias_interface import BiasInterface, FileLoaderThread


def testBiasDocumentServiceLoadsExcelWithOriginalRowNumbers(tmp_path) -> None:
    path = tmp_path / "sample.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["作文", "等级"])
    sheet.append(["第一句[C]", "HSK 4"])
    sheet.append(["第二句{CC词}", "HSK 5"])
    workbook.save(path)

    dataFrame, rowCount = BiasDocumentService().loadFile(str(path))

    assert rowCount == 2
    assert list(dataFrame.columns) == ["作文", "等级"]
    assert dataFrame.attrs["sourceKind"] == "excel"
    assert dataFrame.attrs["sourcePositions"] == [2, 3]


def testBiasDocumentServiceLoadsTxtWithEncodingAndLineNumbers(tmp_path) -> None:
    path = tmp_path / "sample.txt"
    path.write_bytes("第一句[C]\n\n第二句{CC词}\n".encode("gb18030"))

    dataFrame, rowCount = BiasDocumentService().loadFile(str(path))

    assert rowCount == 2
    assert list(dataFrame.columns) == [BIAS_TEXT_COLUMN]
    assert dataFrame[BIAS_TEXT_COLUMN].tolist() == ["第一句[C]", "第二句{CC词}"]
    assert dataFrame.attrs["sourceKind"] == "txt"
    assert dataFrame.attrs["sourcePositions"] == [1, 3]


def testBiasDocumentServiceLoadsDocxParagraphsAndTableRows(tmp_path) -> None:
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("第一段[C]")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "第二段"
    table.cell(0, 1).text = "{CC词}"
    document.add_paragraph("")
    document.save(path)

    dataFrame, rowCount = BiasDocumentService().loadFile(str(path))

    assert rowCount == 2
    assert dataFrame[BIAS_TEXT_COLUMN].tolist() == ["第一段[C]", "第二段\t{CC词}"]
    assert dataFrame.attrs["sourceKind"] == "docx"
    assert dataFrame.attrs["sourcePositions"] == [1, 2]


def testBiasDocumentServiceExtractsLegacyDocPieceTable() -> None:
    text = "第一段[C]\r第二段{CC词}"
    encodedText = text.encode("utf-16-le")
    textOffset = 1024
    wordStream = bytearray(textOffset + len(encodedText))
    struct.pack_into("<H", wordStream, 0, 0xA5EC)
    struct.pack_into("<H", wordStream, 10, 0x0200)

    cursor = 32
    struct.pack_into("<H", wordStream, cursor, 0)
    cursor += 2
    struct.pack_into("<H", wordStream, cursor, 22)
    cursor += 2
    struct.pack_into("<I", wordStream, cursor + 12, len(text))
    cursor += 22 * 4
    struct.pack_into("<H", wordStream, cursor, 34)
    cursor += 2

    plcPcd = struct.pack("<IIHIH", 0, len(text), 0, textOffset, 0)
    clx = b"\x02" + struct.pack("<I", len(plcPcd)) + plcPcd
    struct.pack_into("<II", wordStream, cursor + 33 * 8, 0, len(clx))
    wordStream[textOffset:] = encodedText

    extracted = BiasDocumentService()._extractLegacyWordText(
        bytes(wordStream),
        clx,
    )

    assert extracted.splitlines() == ["第一段[C]", "第二段{CC词}"]


def testBiasDocumentServiceRejectsEmptyTextFile(tmp_path) -> None:
    path = tmp_path / "empty.txt"
    path.write_text(" \n\n", encoding="utf-8")

    try:
        BiasDocumentService().loadFile(str(path))
    except BiasDocumentLoadError as error:
        assert "没有可分析的文本" in str(error)
    else:
        raise AssertionError("空文本文件应被拒绝")


def testFileLoaderThreadDelegatesParsingToService() -> None:
    class FakeDocumentService:
        def __init__(self) -> None:
            self.loadedPath = None

        def loadFile(self, filePath: str):
            self.loadedPath = filePath
            return "frame", 3

    service = FakeDocumentService()
    loader = FileLoaderThread(["sample.txt"], documentService=service)

    assert loader._loadFile("sample.txt") == ("frame", 3)
    assert service.loadedPath == "sample.txt"


def testBiasInterfacePresentsAndAcceptsDocumentSources(qtbot, tmp_path) -> None:
    path = tmp_path / "sample.txt"
    path.write_text("第一句[C]\n第二句{CC词}", encoding="utf-8")
    dataFrame, rowCount = BiasDocumentService().loadFile(str(path))

    interface = BiasInterface()
    qtbot.addWidget(interface)
    interface._onFileLoaded(str(path), dataFrame, rowCount)

    assert interface.chooseFileBtn.text() == "选择文件"
    assert "XLSX、TXT、DOCX 和 DOC" in interface.chooseFileBtn.toolTip()
    assert interface.columnCombobox.currentText() == BIAS_TEXT_COLUMN
    assert interface.columnConfigBtn.isEnabled() is False
    assert "2 条文本" in interface.sourceStatusLabel.text()
