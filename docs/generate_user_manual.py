# coding: utf-8
"""
Prismatica（棱溯客户端）用户使用手册生成脚本

使用 python-docx 生成完整的 .docx 用户手册，章节结构如下：
  封面
  目录（自动生成 TOC 字段，Word 打开后按 F9 / 右键更新）
  第 1 章 概述
  第 2 章 安装与启动
  第 3 章 界面总览
  第 4 章 HSK 动态作文语料库下载
  第 5 章 全球中介语语料库下载
  第 6 章 偏误统计
  第 7 章 语料分析（10 个子模块）
  第 8 章 任务管理
  第 9 章 设置
  附录 A 常见问题
  附录 B 版本与版权

排版规范：
  - 页面：A4，上下 2.5cm，左右 2.5cm
  - 正文字体：宋体 / Times New Roman 五号（10.5pt）
  - 标题：黑体 / Arial 黑色加粗
  - 封面单独一节，不显示页码
  - 正文从第 1 页起，页脚显示页码与软件名
"""

from __future__ import annotations

import os
from datetime import date
from typing import Iterable, List, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------
# 常量
# ---------------------------------------------------------------------

APP_NAME = "Prismatica（棱溯客户端）"
APP_VERSION = "v1.0.0"
APP_AUTHOR = "猫叁零"
APP_ORG = "贵州六棱光界科技工作室"
APP_YEAR = 2026

OUTPUT_DIR = r"E:\Prismatica\docs"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, f"Prismatica_用户使用手册_{APP_VERSION}.docx")

# 排版常量
CN_FONT = "宋体"
CN_HEADING_FONT = "黑体"
EN_FONT = "Times New Roman"
EN_HEADING_FONT = "Arial"

# 颜色（参考 qfluentwidgets 主题色 #00b09c）
THEME_COLOR = RGBColor(0x00, 0xB0, 0x9C)
ACCENT_COLOR = RGBColor(0x32, 0x75, 0x9C)
SUBTLE_GRAY = RGBColor(0x88, 0x88, 0x88)


# ---------------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------------

def set_run_font(run, name_cn: str = CN_FONT, name_en: str = EN_FONT,
                 size_pt: float = 10.5, bold: bool = False,
                 color: RGBColor | None = None) -> None:
    """统一设置 run 的中英文字体、字号、颜色。"""
    run.font.name = name_en
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name_en)
    rfonts.set(qn("w:hAnsi"), name_en)
    rfonts.set(qn("w:eastAsia"), name_cn)
    rfonts.set(qn("w:cs"), name_en)


def set_paragraph_spacing(paragraph, before_pt: float = 0, after_pt: float = 4,
                          line_spacing: float | None = 1.5,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """统一段落间距与行距。"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.alignment = alignment


def add_paragraph(doc, text: str, *, size_pt: float = 10.5,
                  bold: bool = False, color: RGBColor | None = None,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  first_line_indent_chars: float = 0,
                  before_pt: float = 0, after_pt: float = 4,
                  line_spacing: float | None = 1.5):
    """添加一个普通正文段落。

    first_line_indent_chars: 首行缩进（中文字符数，2 = 缩进 2 个汉字）。
    """
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=before_pt, after_pt=after_pt,
                          line_spacing=line_spacing, alignment=alignment)
    if first_line_indent_chars:
        p.paragraph_format.first_line_indent = Pt(size_pt * first_line_indent_chars)
    if text:
        run = p.add_run(text)
        set_run_font(run, size_pt=size_pt, bold=bold, color=color)
    return p


def add_heading(doc, text: str, level: int = 1, *, color: RGBColor | None = None):
    """添加自定义标题样式（使用 heading 1/2/3 以便自动进入目录）。"""
    heading_sizes = {1: 18, 2: 14, 3: 12}
    size = heading_sizes.get(level, 11)
    heading = doc.add_heading(level=level)
    # 清理 Word 默认主题色，改用我们自定义颜色
    if color is None:
        color = THEME_COLOR if level == 1 else (
            ACCENT_COLOR if level == 2 else RGBColor(0x33, 0x33, 0x33)
        )
    run = heading.add_run(text)
    name_cn = CN_HEADING_FONT
    name_en = EN_HEADING_FONT
    set_run_font(run, name_cn=name_cn, name_en=name_en, size_pt=size,
                 bold=True, color=color)
    # 标题段前/段后
    pf = heading.paragraph_format
    if level == 1:
        pf.space_before = Pt(18)
        pf.space_after = Pt(10)
    elif level == 2:
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
    else:
        pf.space_before = Pt(8)
        pf.space_after = Pt(4)
    pf.keep_with_next = True
    return heading


def add_bullet_list(doc, items: Iterable[str], *, size_pt: float = 10.5):
    """添加项目符号列表。"""
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, before_pt=0, after_pt=2, line_spacing=1.4)
        # 项目符号与文字间已经有缩进,这里我们简单覆盖
        run = p.add_run(item)
        set_run_font(run, size_pt=size_pt)


def add_numbered_list(doc, items: Iterable[str], *, size_pt: float = 10.5):
    """添加数字列表。"""
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_spacing(p, before_pt=0, after_pt=2, line_spacing=1.4)
        run = p.add_run(item)
        set_run_font(run, size_pt=size_pt)


def add_table(doc, headers: List[str], rows: List[List[str]], *,
              first_col_width_cm: float | None = None,
              total_width_cm: float = 15.0) -> None:
    """添加带表头的表格。

    表头行加底色加粗；表格整体使用浅灰边框。
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # 设置列宽
    col_count = len(headers)
    if first_col_width_cm is not None and col_count > 1:
        first_w = first_col_width_cm
        rest_w = (total_width_cm - first_w) / (col_count - 1)
        widths = [first_w] + [rest_w] * (col_count - 1)
    else:
        widths = [total_width_cm / col_count] * col_count
    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = Cm(w)

    # 表头
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 清空默认段落, 重新设置字体
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        set_run_font(run, size_pt=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        # 表头底色
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "00B09C")
        tc_pr.append(shd)

    # 数据行
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(str(value))
            set_run_font(run, size_pt=10.5)
        for j, w in enumerate(widths):
            cell.width = Cm(w)

    # 表后空行
    add_paragraph(doc, "", after_pt=2)


def add_tip_box(doc, label: str, content: str, *,
                label_color: RGBColor = THEME_COLOR) -> None:
    """添加一个浅色提示框（用 1x1 表格实现）。"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Cm(15)
    # 整格底色
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0FAF8")
    tc_pr.append(shd)
    # 单元格段落
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(para, before_pt=2, after_pt=2, line_spacing=1.4)
    run = para.add_run(f"{label}  ")
    set_run_font(run, size_pt=10.5, bold=True, color=label_color)
    run2 = para.add_run(content)
    set_run_font(run2, size_pt=10.5)
    add_paragraph(doc, "", after_pt=2)


def insert_toc(doc) -> None:
    """插入 Word 原生 TOC 字段，打开后按 F9 即可更新。"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=1.5)
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "请右键此处选择「更新域」以生成目录"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(placeholder)
    run._r.append(fld_char_end)
    set_run_font(run, size_pt=10.5, color=SUBTLE_GRAY)


def add_page_break(doc) -> None:
    """插入分页符（不开启新节）。"""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def setup_page(section, *, header_text: str = "",
               show_page_number: bool = True,
               different_first_page: bool = False) -> None:
    """统一的页面设置：上下 2.5cm / 左右 2.5cm / 页眉页脚。"""
    section.page_width = Cm(21.0)   # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)
    section.different_first_page_header_footer = different_first_page

    # 页眉
    header = section.header
    header.is_linked_to_previous = False
    if header_text:
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in hp.runs:
            r.text = ""
        hr = hp.add_run(header_text)
        set_run_font(hr, size_pt=9, color=SUBTLE_GRAY)
        # 在页眉段落底部加一条浅灰线
        p_pr = hp._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "DDDDDD")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    # 页脚
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in fp.runs:
        r.text = ""
    if show_page_number:
        # "第 X 页 / 共 Y 页" 形式
        run1 = fp.add_run("第 ")
        set_run_font(run1, size_pt=9, color=SUBTLE_GRAY)
        _add_page_field(fp, "PAGE")
        run2 = fp.add_run(" 页 / 共 ")
        set_run_font(run2, size_pt=9, color=SUBTLE_GRAY)
        _add_page_field(fp, "NUMPAGES")
        run3 = fp.add_run(" 页")
        set_run_font(run3, size_pt=9, color=SUBTLE_GRAY)
    else:
        # 封面页脚：仅一行版权
        run = fp.add_run(f"© {APP_YEAR} {APP_ORG}")
        set_run_font(run, size_pt=9, color=SUBTLE_GRAY)


def _add_page_field(paragraph, field_name: str) -> None:
    """在段落里追加一个 PAGE / NUMPAGES 字段。"""
    run = paragraph.add_run()
    set_run_font(run, size_pt=9, color=SUBTLE_GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_name} "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


# ---------------------------------------------------------------------
# 章节内容生成
# ---------------------------------------------------------------------

def build_cover(doc: Document) -> None:
    """封面页（不显示页码）。"""
    # 顶部空 6 行
    for _ in range(6):
        add_paragraph(doc, "", after_pt=0)

    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_pt=0, after_pt=12, line_spacing=1.2)
    run = p.add_run("Prismatica")
    set_run_font(run, name_cn="Arial", name_en="Arial",
                 size_pt=44, bold=True, color=THEME_COLOR)

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_pt=0, after_pt=4, line_spacing=1.2)
    run = p.add_run("棱溯客户端")
    set_run_font(run, name_cn=CN_HEADING_FONT, name_en=EN_HEADING_FONT,
                 size_pt=28, bold=True, color=ACCENT_COLOR)

    # 中文文档名
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_pt=10, after_pt=4, line_spacing=1.2)
    run = p.add_run("用户使用手册")
    set_run_font(run, name_cn=CN_HEADING_FONT, name_en=EN_HEADING_FONT,
                 size_pt=22, bold=True, color=RGBColor(0x33, 0x33, 0x33))

    # 一行装饰横线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_pt=10, after_pt=10, line_spacing=1.0)
    run = p.add_run("━" * 20)
    set_run_font(run, name_cn="Arial", name_en="Arial",
                 size_pt=14, color=THEME_COLOR)

    # Slogan
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_pt=20, after_pt=4, line_spacing=1.5)
    run = p.add_run("中文学术语料处理一站式桌面应用")
    set_run_font(run, size_pt=14, color=RGBColor(0x66, 0x66, 0x66))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_pt=0, after_pt=4, line_spacing=1.5)
    run = p.add_run("对标 AntConc · 面向中文教学与研究")
    set_run_font(run, size_pt=12, color=SUBTLE_GRAY)

    # 中部留白
    for _ in range(8):
        add_paragraph(doc, "", after_pt=0)

    # 底部元信息
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    for i, (k, v) in enumerate([
        ("版本号", APP_VERSION),
        ("适用对象", "中文学术语料研究者、汉语教师及研究生"),
        ("编写方", APP_ORG),
        ("发布日期", date.today().strftime("%Y 年 %m 月 %d 日")),
    ]):
        row = info_table.rows[i]
        for cell, w in zip(row.cells, [Cm(4.0), Cm(10.0)]):
            cell.width = w
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 键
        kc = row.cells[0]
        for p in kc.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in p.runs:
                r.text = ""
        kp = kc.paragraphs[0]
        kp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        kr = kp.add_run(k)
        set_run_font(kr, size_pt=11, bold=True, color=SUBTLE_GRAY)
        # 值
        vc = row.cells[1]
        for p in vc.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.text = ""
        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vr = vp.add_run("    " + v)
        set_run_font(vr, size_pt=11)


def build_toc(doc: Document) -> None:
    """目录页。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_pt=0, after_pt=12, line_spacing=1.2)
    run = p.add_run("目  录")
    set_run_font(run, name_cn=CN_HEADING_FONT, name_en=EN_HEADING_FONT,
                 size_pt=20, bold=True, color=THEME_COLOR)
    insert_toc(doc)
    add_paragraph(doc, "", after_pt=2)
    add_tip_box(doc, "提示",
                "在 Microsoft Word 中打开本文档后，"
                "请右键上方目录区域选择「更新域」→「更新整个目录」"
                "（或直接按 F9），即可自动生成完整目录与页码。")


# ---------------------------------------------------------------------
# 正文章节
# ---------------------------------------------------------------------

def chapter_1_overview(doc: Document) -> None:
    add_heading(doc, "第 1 章  概述", level=1)

    add_heading(doc, "1.1  软件简介", level=2)
    add_paragraph(doc,
        "Prismatica（中文名「棱溯客户端」）是一款面向中文学术语料处理研究的桌面应用，"
        "对标国际语料库语言学经典工具 AntConc，针对中文场景进行了深度适配。"
        "软件集成了语料下载、偏误统计、词频分析、KWIC 检索、搭配抽取、"
        "共现网络、句法依存、情感分析、词云可视化等核心能力，"
        "并以图形化界面屏蔽底层命令，"
        "让汉语教师、中文研究方向的学生与研究者能够聚焦于语言现象本身。",
        first_line_indent_chars=2)

    add_paragraph(doc,
        "本软件基于 Python 3.11 与 PySide6 构建，UI 部分使用 qfluentwidgets Pro 主题框架，"
        "整体风格统一、现代，符合 Windows 11 Fluent Design 设计语言。",
        first_line_indent_chars=2)

    add_heading(doc, "1.2  主要功能一览", level=2)
    add_table(doc,
        headers=["模块", "对应导航", "核心能力"],
        rows=[
            ["HSK 语料下载", "HSK下载", "HSK 动态作文语料库的四种检索方式与后台下载任务管理"],
            ["全球中介语下载", "全球中介下载", "北京语言大学全球中介语语料库的字符串/条件/搭配/词性检索"],
            ["偏误统计", "偏误统计", "从 Excel 作文标注中自动识别 33 类偏误，生成计数/图表/热力图/关联规则"],
            ["语料分析", "语料分析", "10 个子模块：语料导入、词频、词语、主题词、KWIC、情感、搭配、词云、网络、依存"],
            ["任务管理", "任务管理", "统一查看进行中与已完成的下载/导入任务，支持取消、重试、定位文件"],
            ["设置", "设置", "下载路径/线程/重试参数、HSK & Global 登录 Token、激活码管理、系统信息"],
        ],
        first_col_width_cm=3.5)

    add_heading(doc, "1.3  软件版本", level=2)
    add_table(doc,
        headers=["字段", "内容"],
        rows=[
            ["产品名称", "Prismatica"],
            ["中文名", "棱溯客户端"],
            ["当前版本", APP_VERSION],
            ["内部版本", "100"],
            ["开发方", APP_ORG],
            ["主要作者", APP_AUTHOR],
            ["首发年份", str(APP_YEAR)],
            ["授权模式", "内测免费 / 正式版激活码授权"],
        ],
        first_col_width_cm=3.5)

    add_heading(doc, "1.4  适用人群", level=2)
    add_bullet_list(doc, [
        "汉语国际教育方向的教师与研究者，需要快速提取偏误特征并可视化呈现。",
        "中文系本科生 / 研究生，需要完成语料库语言学课程作业或学位论文。",
        "面向中文的 NLP 研究者，需要对小型语料做统计分析与可视化对比。",
        "HSK / 全球中介语语料库的日常使用者，希望摆脱网页检索的繁琐操作。",
    ])


def chapter_2_install(doc: Document) -> None:
    add_heading(doc, "第 2 章  安装与启动", level=1)

    add_heading(doc, "2.1  系统要求", level=2)
    add_table(doc,
        headers=["项目", "最低要求", "推荐配置"],
        rows=[
            ["操作系统", "Windows 10 (1809+)", "Windows 11 22H2+"],
            ["CPU", "x86_64 4 核", "x86_64 6 核及以上"],
            ["内存", "8 GB", "16 GB（处理 100 万字以上语料时）"],
            ["磁盘", "可用空间 5 GB", "SSD，可用空间 20 GB"],
            ["Python", "3.11.x（仅源码运行需要）", "3.11.x"],
            ["网络", "下载语料时需要联网", "稳定的宽带连接"],
        ],
        first_col_width_cm=3.0)

    add_heading(doc, "2.2  启动方式", level=2)

    add_heading(doc, "2.2.1  源码运行（开发模式）", level=3)
    add_paragraph(doc, "项目使用 uv 管理依赖，环境配置步骤如下：",
                  first_line_indent_chars=2)
    add_numbered_list(doc, [
        "安装 uv：Windows 推荐使用 pip install uv，或前往 astral.sh 下载安装包。",
        "在项目根目录执行 uv sync，自动创建 .venv 并安装全部依赖。",
        f"执行 uv run python main.py 启动 GUI。",
        "（可选）VS Code 调试：按 F5 启动，配置位于 .vscode/launch.json。",
    ])

    add_heading(doc, "2.2.2  打包版本（EXE）", level=3)
    add_paragraph(doc,
        "正式版以单文件 EXE 形式发布。直接双击 Prismatica.exe 即可启动，"
        "首次启动会创建数据目录（详见 2.3 节）。如弹出 Windows Defender 警告，"
        "请选择「仍要运行」（签名证书后续会更新）。",
        first_line_indent_chars=2)

    add_heading(doc, "2.3  数据目录结构", level=2)
    add_paragraph(doc,
        "软件启动时会在安装目录（即 EXE 所在目录，或开发模式下的项目根）"
        "下自动创建以下目录：", first_line_indent_chars=2)
    add_table(doc,
        headers=["目录", "用途", "是否随项目仓库提交"],
        rows=[
            ["config/", "软件配置文件（config.json、用户自定义清洗预设等）", "否（已 gitignore）"],
            ["download/", "通过 HSK / Global 下载的原始语料文件", "否"],
            ["logs/", "应用运行日志，按日期轮转", "否"],
            ["datas/", "语料库注册表、各语料库独立 SQLite、导出文件", "否"],
            ["datas/corpora/", "每个语料库一个 .db 文件（SQLite + FTS5）", "否"],
            ["datas/exports/", "导出的报告、图表、CSV 文件", "否"],
        ],
        first_col_width_cm=3.5)
    add_tip_box(doc, "提示",
                "软件升级或重装时，datas/、config/、download/ 目录里的数据会自动保留。"
                "如需彻底重置，可直接删除 datas/ 与 config/ 目录。")

    add_heading(doc, "2.4  首次启动", level=2)
    add_numbered_list(doc, [
        "双击 EXE 或运行 main.py，等待 2~3 秒出现启动屏（Splash Screen）。",
        "软件自动初始化：注入 qfluentwidgets Pro 授权 → 配置日志 → "
        "迁移旧版数据 → 加载语料库注册表 → 实例化主窗口。",
        "默认进入「HSK 下载」页面，左侧导航栏共 6 个模块。",
        "建议进入「设置」页确认下载保存路径、线程数、Token 是否正确。",
    ])

    add_heading(doc, "2.5  升级与降级", level=2)
    add_paragraph(doc,
        "升级：新版本直接覆盖 EXE 即可，datas/ 与 config/ 目录不动，"
        "注册表与已导入语料会完整保留。", first_line_indent_chars=2)
    add_paragraph(doc,
        "降级：不建议。如确需回退旧版，请先在「设置 → 下载功能设置」"
        "中备份自定义配置，再覆盖 EXE，最后回拷配置。", first_line_indent_chars=2)


def chapter_3_ui_overview(doc: Document) -> None:
    add_heading(doc, "第 3 章  界面总览", level=1)

    add_heading(doc, "3.1  主窗口布局", level=2)
    add_paragraph(doc,
        "主窗口采用 Fluent Design 设计语言，整体分为四个区域：",
        first_line_indent_chars=2)
    add_bullet_list(doc, [
        "标题栏（顶部）：显示软件 Logo 与名称「棱溯客户端」。",
        "导航栏（左侧）：6 个模块入口，顶部 4 个为功能页（HSK / Global / 偏误 / 语料分析），"
        "底部 2 个为辅助页（任务管理 / 设置）。",
        "工作区（中间）：当前导航对应页面的内容。",
        "状态栏：右下角 InfoBar 用于显示任务完成、错误等瞬时提示。",
    ])

    add_heading(doc, "3.2  导航栏模块对照", level=2)
    add_table(doc,
        headers=["位置", "图标", "模块", "一句话简介"],
        rows=[
            ["顶部", "📚 HSK", "HSK 下载", "HSK 动态作文语料库检索与下载"],
            ["顶部", "🌐 Global", "全球中介下载", "北京语言大学全球中介语语料库检索与下载"],
            ["顶部", "📊 Bias", "偏误统计", "Excel 偏误标注的统计与可视化"],
            ["顶部", "🔍 Analysis", "语料分析", "对标 AntConc 的 10 个子分析模块"],
            ["底部", "📋 Task", "任务管理", "下载/导入任务进度跟踪"],
            ["底部", "⚙️ Setting", "设置", "下载参数、Token、激活码管理"],
        ],
        first_col_width_cm=2.0)

    add_heading(doc, "3.3  通用交互约定", level=2)
    add_bullet_list(doc, [
        "InfoBar 提示：所有成功/失败/警告反馈通过右上角 InfoBar 显示，停留 2~3 秒自动消失。",
        "后台线程：耗时操作（下载、词频计算、网络图构建等）均在 QThread 中执行，"
        "UI 全程可响应。",
        "关闭主窗口时若有进行中任务，会弹出确认对话框，可选择「等待完成」或「强制停止」。",
        "数据持久化：所有配置、语料库、Token 均自动保存，无需手动保存。",
    ])

    add_heading(doc, "3.4  主题色与外观", level=2)
    add_paragraph(doc,
        "软件主色为薄荷绿（#00B09C），与 qfluentwidgets 默认主题搭配。"
        "暂未提供主题切换功能，后续会按用户反馈开放深色模式。",
        first_line_indent_chars=2)


def chapter_4_hsk(doc: Document) -> None:
    add_heading(doc, "第 4 章  HSK 动态作文语料库下载", level=1)

    add_heading(doc, "4.1  模块概览", level=2)
    add_paragraph(doc,
        "本模块对接北京语言大学 HSK 动态作文语料库（hsk.blcu.edu.cn），"
        "提供 4 种检索方式，所有检索结果都会以「任务」形式提交到后台下载队列，"
        "下载进度可在「任务管理」页面查看。",
        first_line_indent_chars=2)
    add_tip_box(doc, "使用前提",
                "首次使用前请在「设置 → 下载功能设置」中点击「HSK-Token → 刷新」"
                "登录获取 Token，Token 默认会自动保存。")

    add_heading(doc, "4.2  检索方式", level=2)

    add_heading(doc, "4.2.1  字符串一般检索", level=3)
    add_paragraph(doc, "最简单、最常用的检索方式，输入一个关键词即可。",
                  first_line_indent_chars=2)
    add_table(doc,
        headers=["参数", "说明", "是否必填"],
        rows=[
            ["关键字", "需要查找的字符串，如「图书馆」", "是"],
            ["每页数量", "服务端返回的每页条数（设置中可调）", "—"],
            ["页码", "检索结果页码，从 1 开始", "—"],
        ],
        first_col_width_cm=3.0)

    add_heading(doc, "4.2.2  特定条件检索", level=3)
    add_paragraph(doc,
        "通过「首字符串 + 前词 + 距离 + 后词 + 尾字符串」组合检索，"
        "支持精确的上下文定位。", first_line_indent_chars=2)
    add_table(doc,
        headers=["参数", "说明"],
        rows=[
            ["首字符串", "匹配句首的字符"],
            ["前词", "关键词前一个词"],
            ["距离", "前后词之间的最大距离（0~100）"],
            ["后词", "关键词后一个词"],
            ["尾字符串", "匹配句尾的字符"],
        ],
        first_col_width_cm=3.0)
    add_tip_box(doc, "注意",
                "至少需要填写 1 个条件才可提交任务。")

    add_heading(doc, "4.2.3  词语搭配检索", level=3)
    add_paragraph(doc,
        "按句法关系检索搭配词对，支持 8 种句法结构：主谓、联合、定中、状中、"
        "述宾、述补、数量、介宾。", first_line_indent_chars=2)
    add_table(doc,
        headers=["参数", "说明"],
        rows=[
            ["关键字或词", "需要查找搭配的关键词"],
            ["句法结构", "下拉选择 8 种句法关系之一"],
        ],
        first_col_width_cm=3.0)

    add_heading(doc, "4.2.4  错句检索", level=3)
    add_paragraph(doc,
        "直接按偏误句式代码检索，下拉菜单覆盖 33 种错句类型，"
        "包括把字句错误、被字句错误、句式杂糅、语序错误等。",
        first_line_indent_chars=2)
    add_paragraph(doc, "常用句式代码（部分）：", first_line_indent_chars=2)
    add_table(doc,
        headers=["代码", "含义"],
        rows=[
            ["CJba", "把字句错误"],
            ["CJbei", "被字句错误"],
            ["CJbi", "比字句错误"],
            ["CJX", "语序错误"],
            ["CJZR", "句式杂糅"],
            ["WWJ", "未完句"],
        ],
        first_col_width_cm=3.0)

    add_heading(doc, "4.3  高级设置（筛选条件）", level=2)
    add_paragraph(doc,
        "无论选择哪种检索方式，都可以在「高级设置」中叠加以下筛选维度，"
        "帮助缩小结果范围：", first_line_indent_chars=2)
    add_table(doc,
        headers=["字段", "可选范围", "说明"],
        rows=[
            ["国籍", "100+ 国家", "包括「不限」选项"],
            ["HSK 等级", "3 / 4 / 5 / 6 等", "留学生参加考试时的等级"],
            ["作文题目", "自由文本", "按题目关键字筛选"],
            ["分数区间", "0~100", "按评分区间筛选"],
            ["写作文体", "记叙 / 议论 / 应用 / 说明", "—"],
            ["作文长度", "短 / 中 / 长", "按字数区间划分"],
            ["性别", "男 / 女", "—"],
            ["第一语言", "母语列表", "作者的第一语言"],
            ["第二语言", "母语列表", "作者的母语之外的语言"],
        ],
        first_col_width_cm=3.0)

    add_heading(doc, "4.4  提交与下载流程", level=2)
    add_numbered_list(doc, [
        "在「高级设置」中按需设置筛选条件。",
        "点击右下角「申请任务」按钮。",
        "弹出下载确认框（DownloadApplyWidget），显示参数总览。",
        "确认无误后点击「确定」，任务进入 TaskManager 队列。",
        "右上角 InfoBar 提示「任务已创建」。",
        "切换到「任务管理」页，查看「进行中」/「已完成」任务卡片。",
        "任务完成后，原始语料文件保存在「设置 → 下载保存路径」目录下。",
    ])

    add_heading(doc, "4.5  常见错误", level=2)
    add_table(doc,
        headers=["症状", "可能原因", "处理方式"],
        rows=[
            ["提示「请输入关键字」", "关键字为空", "填写至少 1 个检索条件"],
            ["提示「请至少填写一个检索条件」", "特定条件检索全部为空", "在 5 个条件中至少填 1 个"],
            ["Token 刷新失败", "账号密码错误 / 网络问题", "在「设置 → HSK-Token」重新登录"],
            ["下载结果为空", "筛选条件过严 / 无匹配", "放宽条件后重试"],
        ],
        first_col_width_cm=3.0)


def chapter_5_global(doc: Document) -> None:
    add_heading(doc, "第 5 章  全球中介语语料库下载", level=1)

    add_heading(doc, "5.1  模块概览", level=2)
    add_paragraph(doc,
        "本模块对接北京语言大学全球中介语语料库（qqk.blcu.edu.cn），"
        "提供 4 种检索方式。与 HSK 模块的区别在于：",
        first_line_indent_chars=2)
    add_bullet_list(doc, [
        "支持按「语料类型」筛选（汉字、词语、句型、段落等）。",
        "支持按词性（名词、动词、形容词、副词、介词、连词、助词、数词、量词、代词等）检索。",
        "词语搭配检索支持「左/右」方向选择与「检索后字符数」控制。",
    ])
    add_tip_box(doc, "使用前提",
                "首次使用前请在「设置 → Global-Token → 刷新」登录获取 Token。")

    add_heading(doc, "5.2  检索方式", level=2)

    add_heading(doc, "5.2.1  字符串一般检索", level=3)
    add_paragraph(doc,
        "输入关键字，在指定语料类型表中查找包含该字符串的样本。",
        first_line_indent_chars=2)

    add_heading(doc, "5.2.2  特定条件检索", level=3)
    add_table(doc,
        headers=["参数", "说明"],
        rows=[
            ["语料类型", "下拉选择：汉字、词语、句型、段落等"],
            ["首字符串", "样本首部匹配串"],
            ["前词", "关键词前一个词"],
            ["距离", "前后词最大距离"],
            ["后词", "关键词后一个词"],
            ["尾字符串", "样本尾部匹配串"],
        ],
        first_col_width_cm=3.0)

    add_heading(doc, "5.2.3  词语搭配检索", level=3)
    add_table(doc,
        headers=["参数", "说明"],
        rows=[
            ["关键字或词", "中心词"],
            ["排序方向", "左 / 右（搭配词相对于中心词的位置）"],
            ["检索后字符数", "返回搭配词的字符长度"],
        ],
        first_col_width_cm=3.0)

    add_heading(doc, "5.2.4  按词性检索", level=3)
    add_paragraph(doc,
        "直接输入词性代码（n 名词、v 动词、a 形容词、d 副词、p 介词、"
        "c 连词、u 助词、m 数词、q 量词、r 代词 等），"
        "检索指定词性在语料中的全部出现。", first_line_indent_chars=2)

    add_heading(doc, "5.3  高级设置", level=2)
    add_table(doc,
        headers=["字段", "说明"],
        rows=[
            ["母语", "作者母语，列表形式"],
            ["HSK 等级", "作者汉语水平等级"],
            ["汉语水平", "作者汉语能力描述"],
            ["作者国籍", "作者所属国家"],
            ["标签", "扩展标签，按服务端支持的标签填写"],
            ["机构 ID", "高级筛选，可不填"],
        ],
        first_col_width_cm=3.0)

    add_heading(doc, "5.4  提交流程", level=2)
    add_paragraph(doc, "整体流程与 HSK 模块一致：填写检索条件 → 点击「申请任务」"
                      "→ 确认参数 → 任务进入队列 → 在「任务管理」查看进度。"
                      "完成后文件保存在「设置 → 下载保存路径」。",
                      first_line_indent_chars=2)


def chapter_6_bias(doc: Document) -> None:
    add_heading(doc, "第 6 章  偏误统计", level=1)

    add_heading(doc, "6.1  模块概览", level=2)
    add_paragraph(doc,
        "本模块用于分析带偏误标注的语料 Excel 文件（典型来源：HSK 动态作文语料库导出）。"
        "支持自动识别 33 类偏误，生成统计表格、饼图/条形图、"
        "按等级/国籍的热力图，以及 Apriori 关联规则挖掘。",
        first_line_indent_chars=2)

    add_heading(doc, "6.2  文件准备", level=2)
    add_paragraph(doc, "支持以下两种 Excel 结构：", first_line_indent_chars=2)
    add_bullet_list(doc, [
        "单文件：每行一句作文或一句标注，列为字段（作文正文、偏误代码、等级、国籍等）。",
        "多文件批量：将多个 Excel 一次性拖入界面，软件会逐个处理。",
    ])
    add_tip_box(doc, "性能提示",
                "超过 10 万行时建议拆分为多个文件，软件采用后台线程加载，"
                "不会阻塞 UI。")

    add_heading(doc, "6.3  偏误代码字典", level=2)
    add_paragraph(doc, "软件内置 33 类偏误代码，分布如下：", first_line_indent_chars=2)

    add_heading(doc, "6.3.1  字形类（10 种）", level=3)
    add_table(doc,
        headers=["代码", "类型", "示例"],
        rows=[
            ["[C]", "错字", "把 /[C]"],
            ["[Bxxx]", "别字", "高兴[B高]"],
            ["[L]", "漏字", "高兴 /[L]"],
            ["[Dxxx]", "多字", "[D了]"],
            ["[Fxxx]", "繁体字", "學[F学]"],
            ["[Yxxx]", "异体字", "羣[Y群]"],
            ["[Pxxx]", "拼音字", "[Pnihao]"],
            ["[#]", "无法识别", "无法识别[#]"],
            ["[BCxxx]", "错误标点", "[BC，]"],
            ["[BQxxx]", "空缺标点", "[BQ。]"],
            ["[BDxxx]", "多余标点", "[BD、]"],
        ],
        first_col_width_cm=2.5)

    add_heading(doc, "6.3.2  句式类（19 种）", level=3)
    add_table(doc,
        headers=["代码", "类型"],
        rows=[
            ["{CJba}", "把字句错误"],
            ["{CJbei}", "被字句错误"],
            ["{CJbi}", "比字句错误"],
            ["{CJl}", "连字句错误"],
            ["{CJy}", "有字句错误"],
            ["{CJs}", "是字句错误"],
            ["{CJsd}", "「是……的」句错误"],
            ["{CJcx}", "存现句错误"],
            ["{CJjy}", "兼语句错误"],
            ["{CJld}", "连动句错误"],
            ["{CJshb}", "双宾语句错误"],
            ["{CJxw}", "形容词谓语句错误"],
            ["{CJ-/+}", "句子成分残缺/多余"],
            ["{CJX}", "语序错误"],
            ["{CJZR}", "句式杂糅"],
            ["{CJcd}", "重叠错误"],
            ["{CJgd}", "固定格式错误"],
            ["{CJ?}", "句处理存疑"],
            ["{WWJ}", "未完句标记"],
        ],
        first_col_width_cm=3.0)

    add_heading(doc, "6.3.3  词语类（6 种）", level=3)
    add_table(doc,
        headers=["代码", "类型"],
        rows=[
            ["{CC}", "错词"],
            ["{CLH}", "离合词错误"],
            ["{W}", "外文词"],
            ["{CQ}", "缺词"],
            ["{CD}", "多词"],
            ["{CY}", "存疑词"],
        ],
        first_col_width_cm=3.0)

    add_heading(doc, "6.4  操作流程", level=2)
    add_numbered_list(doc, [
        "点击「选择文件」或拖入 Excel 文件，支持批量。",
        "等待后台线程加载完成（顶部进度条显示文件名）。",
        "在「偏误类型」区域通过多选框勾选要分析的类型；"
        "点击「全选/取消」可一键切换。",
        "点击「开始分析」按钮，软件按选定类型统计。",
        "分析完成后，结果区域展示：",
    ])
    add_bullet_list(doc, [
        "类型计数表：每类偏误的条数与占比，支持排序。",
        "饼图 / 条形图：可视化偏误分布，可导出 PNG / SVG 或复制到剪贴板。",
        "热力图：按 HSK 等级或作者国籍展示偏误分布，支持点击单元格下钻。",
        "关联规则（Apriori）：挖掘偏误类型之间的共现模式，"
        "支持最小支持度 / 最小置信度参数调整，散点图与网络图双视图。",
    ])

    add_heading(doc, "6.5  列配置（等级 / 国籍）", level=2)
    add_paragraph(doc,
        "热力图分组依赖「等级」与「国籍」列。点击工具栏「列配置」按钮，"
        "在弹窗中选择对应的 Excel 表头：", first_line_indent_chars=2)
    add_bullet_list(doc, [
        "「等级」列：包含 level / hsk / 等级 / 级别 / 水准 等关键词时会被自动识别。",
        "「国籍」列：包含 country / nationality / 国籍 / 国家 / nation 等关键词时被自动识别。",
        "点击「根据列名自动识别」按钮，软件会按上述规则自动填充。",
    ])

    add_heading(doc, "6.6  关联规则（Apriori）", level=2)
    add_paragraph(doc,
        "对每篇作文中出现的偏误类型组合进行事务编码，"
        "使用 mlxtend 库提供的 Apriori 算法挖掘频繁项集与关联规则。"
        "结果展示三种视图：", first_line_indent_chars=2)
    add_bullet_list(doc, [
        "表格视图：列出每条规则的前项、后项、支持度、置信度、提升度、杠杆值，"
        "可按任意列排序。",
        "散点图视图：以支持度为 x 轴、置信度为 y 轴绘制，支持度-置信度越高表示规则越强。",
        "网络图视图：将规则以前项→后项的有向图形式展示，"
        "边粗细映射置信度，节点大小映射频次。",
    ])
    add_tip_box(doc, "阈值建议",
                "对于 1 万条以内的偏误样本，建议将「最小支持度」设为 0.05~0.10，"
                "「最小置信度」设为 0.5~0.7。样本量很小时可适当降低。")

    add_heading(doc, "6.7  导出", level=2)
    add_table(doc,
        headers=["导出项", "支持格式", "保存位置"],
        rows=[
            ["图表（饼图/条形图/热力图）", "PNG / SVG", "用户选择"],
            ["关联规则表", "CSV", "用户选择"],
            ["统计结果", "复制到剪贴板", "—"],
        ],
        first_col_width_cm=4.0)


def chapter_7_corpus_analysis(doc: Document) -> None:
    add_heading(doc, "第 7 章  语料分析", level=1)

    add_heading(doc, "7.1  模块概览", level=2)
    add_paragraph(doc,
        "本模块是 Prismatica 的核心，包含 10 个子分析面板，"
        "对标 AntConc 的核心功能并针对中文场景进行了扩展。"
        "页面顶部为分段控件（Segmented Widget），点击标签切换不同分析视图。",
        first_line_indent_chars=2)
    add_paragraph(doc,
        "所有面板共享同一份「当前语料库」：通过顶部语料库切换器（CorpusSwitcherWidget）"
        "可在多个已注册语料库之间切换。", first_line_indent_chars=2)

    add_table(doc,
        headers=["子模块", "对标 AntConc", "一句话功能"],
        rows=[
            ["语料导入与清洗", "—", "Excel/TXT/DOCX 导入与文本清洗规则配置"],
            ["词频分析", "Word List", "单频词与 N-gram 词频统计"],
            ["词语分析", "Word List + 词频分布", "词汇指标、增长曲线、高频词、词汇分布"],
            ["主题词分析", "Keyword List", "对照参照语料库找出过度出现的关键词"],
            ["语境分析", "Concordance (KWIC)", "关键词居中的上下文检索"],
            ["情感分析", "—", "三级情感分类（篇章/段落/句子）与情感分布图"],
            ["搭配分析", "Collocates", "MI / MI3 / T / LogDice / Z / Delta-P 六大搭配强度"],
            ["词语云图", "Word Cluster", "可定制形状/配色/字体的词云图"],
            ["共现网络图", "—", "滑动窗口共现 + 力导向布局 + 社区发现"],
            ["句法依存图", "—", "HanLP 句法依存可视化（树状/弧状）"],
        ],
        first_col_width_cm=3.5)

    add_heading(doc, "7.2  语料导入与清洗", level=2)
    add_paragraph(doc, "首个面板，承担「数据准备」职责：", first_line_indent_chars=2)
    add_numbered_list(doc, [
        "支持格式：.txt（自动嗅探 utf-8/gbk/utf-16/latin-1）、.md、.docx、.xlsx。",
        "批量导入：将多个文件一次性拖入；后台线程流式读取，不阻塞 UI。",
        "语料库管理：通过「语料库切换器」创建 / 重命名 / 删除语料库。"
        "每个语料库对应 datas/corpora/ 下的独立 SQLite 文件。",
        "清洗规则：支持自定义正则替换、停用词表、自定义词典。"
        "内置与「用户自定义」双目录预设，预设文件为 JSON 格式。",
        "分词缓存：分词结果会缓存到 TokenCache（带模型版本号与文本哈希），"
        "模型升级时自动失效。",
    ])

    add_heading(doc, "7.3  词频分析", level=2)
    add_table(doc,
        headers=["参数", "说明", "可选范围"],
        rows=[
            ["最小词长", "过滤短词的字符数阈值", "1~10"],
            ["最大词长", "过滤长词的字符数阈值", "1~50"],
            ["大小写敏感", "是否区分大小写", "开 / 关"],
            ["过滤数字", "是否排除纯数字 token", "开 / 关"],
            ["启用停用词", "使用停用词表过滤", "开 / 关"],
            ["使用 jieba 分词", "中文分词", "开 / 关（关闭则按字符切分）"],
            ["N-gram 阶数", "Bigram / Trigram / ...", "2~5"],
            ["N-gram 最低频次", "过滤低频 n-gram", "1~"],
        ],
        first_col_width_cm=3.5)
    add_paragraph(doc,
        "结果以表格形式展示（词、频次、累计频次、占比），"
        "可一键导出 CSV、绘制 Zipf 分布图、查看 N-gram 聚类。",
        first_line_indent_chars=2)

    add_heading(doc, "7.4  词语分析", level=2)
    add_bullet_list(doc, [
        "词汇指标卡：词数（Type）、形符数（Token）、"
        "类符/形符比（TTR）、Guiraud 指数、Herdan 指数、Uber 指数、平均词长。",
        "词汇增长曲线：Type-Token Curve 展示样本量扩大时新词增长趋势。",
        "高频词列表：含累计 % 与 50% / 80% / 90% 覆盖率标记。",
        "词汇分布：按子库或文件统计词汇分布，支持横向对比。",
    ])

    add_heading(doc, "7.5  主题词分析（Keyness）", level=2)
    add_paragraph(doc,
        "对标 AntConc 的 Keyword List 功能：用户加载「参照语料库 + 观察语料库」"
        "两个语料，软件使用对数似然比（Log-Likelihood, LL）"
        "计算每个词的关键性，输出按 LL 值排序的关键词表。",
        first_line_indent_chars=2)
    add_paragraph(doc, "操作步骤：", first_line_indent_chars=2)
    add_numbered_list(doc, [
        "确保已经创建至少 2 个语料库（观察 + 参照）。",
        "在「主题词分析」面板，通过顶部「语料库切换器」分别指定观察语料与参照语料。",
        "点击「开始计算」，后台线程执行分词 + LL 计算。",
        "结果包括：关键词、观察频次、参照频次、LL 值、p 值、效应量。",
        "可导出 CSV 与图表。",
    ])

    add_heading(doc, "7.6  语境分析（KWIC）", level=2)
    add_table(doc,
        headers=["参数", "说明"],
        rows=[
            ["关键词", "高亮显示在表格中间列"],
            ["左语境宽度", "关键词左侧显示的字符数"],
            ["右语境宽度", "关键词右侧显示的字符数"],
            ["排序方式", "左 1 词 / 左 2 词 / 右 1 词 / 右 2 词"],
            ["二次检索", "在已有 KWIC 结果上嵌套新关键词"],
            ["随机抽样", "指定行数随机抽取"],
            ["上下文扩展", "详情弹窗查看完整句子"],
        ],
        first_col_width_cm=3.5)
    add_paragraph(doc, "结果支持导出 TXT 与 CSV。点击「分布图」按钮"
                      "可生成 KWIC 节点词的左右分布直方图。",
                      first_line_indent_chars=2)

    add_heading(doc, "7.7  情感分析", level=2)
    add_bullet_list(doc, [
        "三级粒度：篇章级、段落级、句子级情感分类（积极/中性/消极）。",
        "情感分布可视化：饼图 + 柱状图双视图。",
        "情感词条：Top N 情感词列表（按出现频次排序）。",
        "自定义情感词典：支持导入用户词典（TXT/JSON）。",
        "报告导出：TXT 格式情感分析报告。",
    ])

    add_heading(doc, "7.8  搭配分析", level=2)
    add_paragraph(doc, "对标 AntConc 的 Collocates 功能，支持 6 种搭配强度指标：",
                  first_line_indent_chars=2)
    add_table(doc,
        headers=["指标", "适用场景", "说明"],
        rows=[
            ["MI（互信息）", "低频但强搭配", "值越大，搭配越强"],
            ["MI3（修正 MI）", "低频 + 高频均稳健", "对低频词不偏倚"],
            ["T-score", "高频搭配", "对高频搭配敏感"],
            ["LogDice", "高频/中频稳健", "值越大越强，范围 0~log2(词表大小)"],
            ["Z-score", "整体显著", "绝对值越大越显著"],
            ["Delta-P", "方向性搭配", "正值表示前项预测后项，负值相反"],
        ],
        first_col_width_cm=2.5)
    add_paragraph(doc,
        "可配置跨距（Span）的 L / R 范围（如 ±1、±3、±5），"
        "结果表格支持跨距位置分布、网络图数据输出与 CSV 导出。",
        first_line_indent_chars=2)

    add_heading(doc, "7.9  词语云图", level=2)
    add_table(doc,
        headers=["参数", "说明", "可选值"],
        rows=[
            ["形状", "词云外形", "矩形 / 圆形 / 椭圆"],
            ["配色", "5 种内置配色", "Cool / Warm / Mono / Viridis / Custom"],
            ["最大词数", "展示的 Top N 词条", "50~500"],
            ["字体", "中文/英文混排字体", "系统已安装的字体"],
            ["旋转", "是否允许词条旋转", "固定 / 随机 / 角度范围"],
        ],
        first_col_width_cm=3.0)
    add_paragraph(doc, "支持导出 PNG 与 SVG。", first_line_indent_chars=2)

    add_heading(doc, "7.10  共现网络图", level=2)
    add_paragraph(doc, "使用 networkx 实现力导向布局，步骤：", first_line_indent_chars=2)
    add_numbered_list(doc, [
        "设置滑动窗口（±N 词，默认 ±2）。",
        "设置筛选阈值：最低词频、最低共现频次、Top K、关键词白名单。",
        "点击「构建网络」，后台线程计算共现矩阵。",
        "完成后进入画布，节点大小映射词频，边粗细映射共现频次，颜色映射社区。",
        "支持 Fruchterman-Reingold 与 spring_layout 切换，"
        "支持社区发现（greedy modularity）着色。",
        "鼠标悬停查看节点/边的详细信息，工具栏支持平移/缩放/保存。",
        "导出：PNG / SVG / GEXF / GraphML。",
    ])

    add_heading(doc, "7.11  句法依存图", level=2)
    add_paragraph(doc,
        "对接 HanLP 句法依存接口，输入文本（支持多句，自动切分）后，"
        "软件会逐句生成依存关系并以两种视图展示：",
        first_line_indent_chars=2)
    add_bullet_list(doc, [
        "树状图：以词为节点、依存关系为有向边，箭头从支配词指向从属词。",
        "弧状图：句子主体在底部水平排列，依存关系以弧线连接。",
    ])
    add_paragraph(doc,
        "支持多句切换（顶部下拉框）、节点悬停高亮，"
        "可导出 PNG / SVG / CoNLL-U 格式。",
        first_line_indent_chars=2)

    add_heading(doc, "7.12  跨面板协作", level=2)
    add_bullet_list(doc, [
        "所有面板共享同一份当前语料库，通过顶部「语料库切换器」一键切换。",
        "分词结果通过 TokenCache 共享，避免重复分词（同一份文本相同模型下只切一次）。",
        "在「词语分析」中点击某个词条，可直接跳转到 KWIC 查看语境。",
        "在「词频分析」中双击某行，可跳转到「搭配分析」并自动填入该词。",
    ])


def chapter_8_task(doc: Document) -> None:
    add_heading(doc, "第 8 章  任务管理", level=1)

    add_heading(doc, "8.1  模块概览", level=2)
    add_paragraph(doc,
        "任务管理是下载与导入任务的统一视图。TaskManager 是单例服务，"
        "负责 pending → in_progress → done / failed 状态流转，"
        "并在主窗口关闭时询问用户是否停止所有进行中任务。",
        first_line_indent_chars=2)

    add_heading(doc, "8.2  页面布局", level=2)
    add_paragraph(doc, "页面顶部为 Pivot 切换器，分为两个 Tab：", first_line_indent_chars=2)
    add_bullet_list(doc, [
        "进行中：显示所有 pending、in_progress 状态的任务卡片。",
        "已完成：显示 done、failed、cancelled 状态的历史任务。",
    ])

    add_heading(doc, "8.3  任务卡片字段", level=2)
    add_table(doc,
        headers=["字段", "说明"],
        rows=[
            ["任务 ID", "系统自动生成的 UUID 短码"],
            ["任务类型", "hskDownload / globalDownload / corpusImport / ... "],
            ["创建时间", "用户点击「申请任务」的时刻"],
            ["进度条", "实时显示下载/导入进度（百分比）"],
            ["状态标签", "排队中 / 进行中 / 已完成 / 失败 / 已取消"],
            ["参数摘要", "检索条件的中文化显示"],
            ["操作按钮", "取消（进行中）/ 重试（失败）/ 打开文件夹（完成）"],
        ],
        first_col_width_cm=3.0)

    add_heading(doc, "8.4  典型操作", level=2)
    add_bullet_list(doc, [
        "取消：点击「取消」按钮，任务进入 cancelled 状态，进入「已完成」列表。",
        "重试：失败任务点击「重试」，会基于原参数重新创建任务。",
        "打开文件夹：完成的任务点击「打开文件夹」，跳转到下载保存目录并选中该文件。",
        "批量停止：主窗口关闭时若有进行中任务，"
        "弹出确认框可选择「等待完成」或「立即停止全部」。",
    ])

    add_heading(doc, "8.5  任务持久化", level=2)
    add_paragraph(doc,
        "任务元数据持久化在 SQLite 中（datas/ 下），"
        "软件重启后会从数据库恢复 pending / in_progress 任务，"
        "确保异常关闭后任务不会丢失。",
        first_line_indent_chars=2)


def chapter_9_settings(doc: Document) -> None:
    add_heading(doc, "第 9 章  设置", level=1)

    add_heading(doc, "9.1  模块概览", level=2)
    add_paragraph(doc,
        "设置页为滚动布局（ScrollArea），从上到下分为 4 个卡片：",
        first_line_indent_chars=2)
    add_bullet_list(doc, [
        "下载功能设置（SoftwareSettingWidget）。",
        "激活码管理（LicenseSettingWidget）。",
        "关于软件（AboutSettingWidget）。",
        "用户协议 + 版权信息（页面底部）。",
    ])

    add_heading(doc, "9.2  下载功能设置", level=2)
    add_table(doc,
        headers=["设置项", "可选值", "默认值", "建议值"],
        rows=[
            ["下载保存路径", "任意可写目录", "安装目录/download/", "D 盘独立分区"],
            ["每页检索数量", "10 / 20 / 50 / 100", "10", "100"],
            ["下载线程数", "1~6", "1", "5（视网络而定）"],
            ["最大重试次数", "1~10", "1", "3"],
            ["HSK-Token", "—", "—", "首次点击「刷新」登录获取"],
            ["Global-Token", "—", "—", "首次点击「刷新」登录获取"],
        ],
        first_col_width_cm=3.0)

    add_heading(doc, "9.2.1  HSK-Token 刷新", level=3)
    add_numbered_list(doc, [
        "点击「HSK-Token → 刷新」按钮。",
        "弹出登录对话框，自动填充上次保存的账号（密码不缓存）。",
        "输入邮箱 + 密码，点击「确定」开始刷新。",
        "刷新成功后，Token 保存到 config/config.json，加密存储。",
        "Token 有效期通常为数天，过期后重复上述流程。",
    ])

    add_heading(doc, "9.2.2  Global-Token 刷新", level=3)
    add_paragraph(doc, "流程与 HSK-Token 一致，区别是账号为 UserID 而非邮箱。",
                  first_line_indent_chars=2)

    add_heading(doc, "9.3  激活码管理", level=2)
    add_paragraph(doc,
        "正式版需要激活码才能使用全部功能；内测期间所有功能免费开放。"
        "点击「管理激活码」按钮弹出独立对话框，包含：",
        first_line_indent_chars=2)
    add_bullet_list(doc, [
        "设备码：32 位字符串 + 复制按钮，激活时需提供给管理员。",
        "激活码输入：正式版输入 16 位以上激活码并点击「激活」。",
        "激活状态：显示「未激活 / 已激活 / 内测版剩余 X 天」等。",
    ])
    add_tip_box(doc, "设备码",
                "设备码基于机器硬件信息（CPU、网卡、磁盘序列号等）生成，"
                "重装系统或更换主板后会变化。")

    add_heading(doc, "9.4  关于软件", level=2)
    add_paragraph(doc, "展示软件版本号、版权方、提交反馈入口与系统信息：",
                  first_line_indent_chars=2)
    add_bullet_list(doc, [
        "版本号：「Prismatica v1.0.0 | 2026 - 猫叁零」。",
        "提交反馈：点击后跳转到在线反馈问卷页。",
        "系统信息：操作系统、CPU 型号、内存、磁盘等。",
    ])

    add_heading(doc, "9.5  用户协议与定价", level=2)
    add_paragraph(doc,
        "页面底部提供两个超链接：「定价协议」「用户协议」，"
        "点击后在系统默认浏览器中打开。", first_line_indent_chars=2)


def appendix_faq(doc: Document) -> None:
    add_heading(doc, "附录 A  常见问题（FAQ）", level=1)

    add_heading(doc, "A.1  启动与安装", level=2)
    add_paragraph(doc, "Q1：双击 EXE 后无反应？", first_line_indent_chars=0,
                  bold=True)
    add_paragraph(doc,
        "A：检查 Windows Defender 是否拦截；右键 EXE → 属性 → 勾选「解除锁定」。"
        "若仍无效，以管理员身份运行。", first_line_indent_chars=2)
    add_paragraph(doc, "Q2：uv sync 失败？", first_line_indent_chars=0, bold=True)
    add_paragraph(doc,
        "A：通常是网络问题。PyPI 已配置清华源，但 pyside6-fluent-widgets-pro"
        " 不在 PyPI，需通过本地 whl 安装（pyproject.toml 已配置）。"
        "若 whl 文件丢失，请从开发方获取。", first_line_indent_chars=2)

    add_heading(doc, "A.2  语料下载", level=2)
    add_paragraph(doc, "Q3：下载任务一直显示「进行中」不结束？", first_line_indent_chars=0,
                  bold=True)
    add_paragraph(doc,
        "A：检查网络。打开日志目录 logs/ 查看是否有网络超时堆栈。"
        "可调低「下载线程数」并增加「最大重试次数」后重试。",
        first_line_indent_chars=2)
    add_paragraph(doc, "Q4：Token 一直刷新失败？", first_line_indent_chars=0, bold=True)
    add_paragraph(doc,
        "A：核对账号密码。注意 HSK 使用邮箱，Global 使用 UserID。"
        "若确认无误仍失败，请检查系统时间是否准确（部分接口会校验时间）。",
        first_line_indent_chars=2)

    add_heading(doc, "A.3  语料分析", level=2)
    add_paragraph(doc, "Q5：词频分析卡顿？", first_line_indent_chars=0, bold=True)
    add_paragraph(doc,
        "A：100 万字以上语料建议关闭 jieba 的新词发现（HSK 词表已足够）。"
        "同时检查是否启用了不必要的停用词表。",
        first_line_indent_chars=2)
    add_paragraph(doc, "Q6：网络图节点太多很乱？", first_line_indent_chars=0, bold=True)
    add_paragraph(doc,
        "A：调高「最低词频」与「最低共现频次」阈值，缩小 Top K。"
        "也可以只保留关键词白名单（手动输入核心词）。",
        first_line_indent_chars=2)
    add_paragraph(doc, "Q7：句法依存返回空？", first_line_indent_chars=0, bold=True)
    add_paragraph(doc,
        "A：检查 HanLP API Key 是否在「设置 → 关于」中正确配置，"
        "或确认网络可访问 HanLP Restful 服务。",
        first_line_indent_chars=2)

    add_heading(doc, "A.4  数据与备份", level=2)
    add_paragraph(doc, "Q8：如何备份语料库与设置？", first_line_indent_chars=0, bold=True)
    add_paragraph(doc,
        "A：复制整个 datas/ 目录与 config/ 目录到外部存储即可。"
        "软件升级前建议做一次完整备份。", first_line_indent_chars=2)
    add_paragraph(doc, "Q9：如何重置所有数据？", first_line_indent_chars=0, bold=True)
    add_paragraph(doc,
        "A：关闭软件后删除 datas/ 与 config/ 目录，重启后会自动重新创建。"
        "此操作不可恢复，请先备份。", first_line_indent_chars=2)

    add_heading(doc, "A.5  错误日志位置", level=2)
    add_paragraph(doc,
        "如需提交 Bug 报告，请附上 logs/ 目录下当天的日志文件，"
        "日志已自动按日期轮转（默认 30 天），并对邮箱、手机号、"
        "身份证号、Token 等敏感信息做了正则脱敏。",
        first_line_indent_chars=2)


def appendix_version(doc: Document) -> None:
    add_heading(doc, "附录 B  版本与版权", level=1)

    add_heading(doc, "B.1  版本历史", level=2)
    add_table(doc,
        headers=["版本", "日期", "变更摘要"],
        rows=[
            ["v1.0.0", "2026-07", "首次正式发布，"
             "包含 HSK / Global 下载、偏误统计、10 个语料分析子模块"],
        ],
        first_col_width_cm=3.0)

    add_heading(doc, "B.2  第三方依赖致谢", level=2)
    add_paragraph(doc, "本软件基于以下开源项目构建（按字母排序）：", first_line_indent_chars=2)
    add_bullet_list(doc, [
        "matplotlib — 图表绘制（PSF / BSD 风格许可证）。",
        "networkx — 图算法（BSD）。",
        "pandas + openpyxl — Excel 读取（BSD / MIT）。",
        "python-docx — Word 文档读取（MIT）。",
        "jieba — 中文分词（MIT）。",
        "HanLP — 句法依存接口调用（Apache 2.0）。",
        "PySide6 + qfluentwidgets — UI 框架（LGPL / 商业授权）。",
        "loguru — 日志（MIT）。",
        "mlxtend — Apriori 算法（BSD）。",
    ])

    add_heading(doc, "B.3  版权声明", level=2)
    add_paragraph(doc,
        f"© {APP_YEAR} {APP_ORG} 保留所有权利。",
        first_line_indent_chars=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        f"{APP_NAME} {APP_VERSION} ｜ 软件著作权登记号：[待补充]",
        first_line_indent_chars=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        "本手册内容如有错漏，欢迎通过「设置 → 提交反馈」告知，"
        "我们会在下一版更新中修正。",
        first_line_indent_chars=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "B.4  联系方式", level=2)
    add_table(doc,
        headers=["渠道", "地址"],
        rows=[
            ["反馈问卷", "https://wj.qq.com/s2/27350075/d71d/"],
            ["定价协议", "https://docs.qq.com/pdf/DTnFzeXhjWXBRd3h0"],
            ["用户协议", "https://docs.qq.com/pdf/DTkhGeXVsWXBGTWN4"],
        ],
        first_col_width_cm=3.0)


# ---------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------

def build_document() -> None:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()

    # ---- 全局默认样式 ----
    style_normal = doc.styles["Normal"]
    style_normal.font.name = EN_FONT
    style_normal.font.size = Pt(10.5)
    rpr = style_normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), EN_FONT)
    rfonts.set(qn("w:hAnsi"), EN_FONT)
    rfonts.set(qn("w:eastAsia"), CN_FONT)

    # =========================================================
    # 第 1 节：封面（不显示页码 + 首页不同）
    # =========================================================
    cover_section = doc.sections[0]
    setup_page(cover_section,
               header_text="",
               show_page_number=False,
               different_first_page=True)
    # 首页不显示页脚
    first_footer = cover_section.first_page_footer
    first_footer.is_linked_to_previous = False
    fp = first_footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in fp.runs:
        r.text = ""
    run = fp.add_run(f"© {APP_YEAR} {APP_ORG}  ·  {APP_NAME} {APP_VERSION}")
    set_run_font(run, size_pt=9, color=SUBTLE_GRAY)

    build_cover(doc)

    # =========================================================
    # 第 2 节：目录（罗马数字页码 I, II, ...）
    # =========================================================
    toc_section = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_page(toc_section,
               header_text=APP_NAME,
               show_page_number=True)
    build_toc(doc)

    # =========================================================
    # 第 3 节：正文（阿拉伯数字页码，从 1 开始）
    # =========================================================
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_page(body_section,
               header_text=APP_NAME + "  ·  用户使用手册",
               show_page_number=True)

    chapter_1_overview(doc)
    chapter_2_install(doc)
    chapter_3_ui_overview(doc)
    chapter_4_hsk(doc)
    chapter_5_global(doc)
    chapter_6_bias(doc)
    chapter_7_corpus_analysis(doc)
    chapter_8_task(doc)
    chapter_9_settings(doc)
    appendix_faq(doc)
    appendix_version(doc)

    # =========================================================
    # 保存
    # =========================================================
    doc.save(OUTPUT_FILE)
    print(f"[OK] 手册已生成: {OUTPUT_FILE}")


if __name__ == "__main__":
    build_document()
